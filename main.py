'''
使用方法:
将需要翻译的word文档存放在src文件夹下，生成的中间文档形式将存放在object目录下
翻译得出的结果将被存放在dst文件中

虚要安装额外的库：
urllib 
execjs
docx

'''



#from translate import Translator #Google提供的API，限制翻译次数，无法大规模翻译
#translator= Translator(from_lang="en",to_lang="zh")
from EN2CH import En2Zh  #用来获取网页反馈的结果
import docx              #用来写入word文件
from win32com import client as wc
from docx import Document
from docx.shared import Inches
import os
import re
src = os.listdir("./src")

base_path = 'D:\python file\调用google实现翻译\src\\'

for item in src:

	document = Document()
	document.add_heading('Translation Results', 0)
	print("add_heading Done!")
	# 首先将doc转换成docx
	word = wc.Dispatch("Word.Application")
	print('docx Done!')
# 找到word路径 + 文件名 ，即可打开文件
	full_path = base_path+item
	doc = word.Documents.Open(full_path)

	# 使用参数16表示将doc转换成docx，保存成docx后才能 读文件
	doc.SaveAs(r"D:\\python file\\调用google实现翻译\\object\\"+item,16)
	print(item,' object Done!')
	doc.Close()
	word.Quit()
	# 读取word内容
	#　这里是以段落为单位的，下面用一个for 遍历所有段落
	doc = docx.Document("D:\python file\调用google实现翻译\\object\\"+item)

	for para in doc.paragraphs:
		all = ''
		sentence = re.split('[.;?!]',para.text)
#		sentence = para.text.split('[.;]')
		for i in sentence:
			item_zh = En2Zh(i).getContent()
#			item_en = translator.translate(i)
			print(item_zh)
			if(item_zh):
				all = all+item_zh+'。'
			print('sentence Done!')
		p = document.add_paragraph(all)
		print('paragraphs Done!')
	document.add_page_break()

	document.save('D:\\python file\\调用google实现翻译\\dst\\'+item)
	print('*********************************')
	print(item,'   destination Done!')
	print('*********************************')
